"""Handler para documentos (OCR, Excel, Word, etc)"""

import os
import logging
from telegram import Update
from telegram.ext import ContextTypes

from security.auth import require_auth
from workspace.core.agent import Agent
from workspace.storage.sqlite_store import SQLiteStore
from config.settings import config

logger = logging.getLogger(__name__)


@require_auth
async def handle_document(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    agent: Agent,
    store: SQLiteStore,
):
    """Handler para documentos (OCR, Excel, Word, etc)"""
    logger.info("Documento recebido")

    await update.message.chat.send_action("typing")

    chat_id = update.effective_chat.id

    try:
        document = update.message.document
        file_name = document.file_name
        mime_type = document.mime_type

        doc_file = await document.get_file()
        doc_path = str(config.TEMP_DIR / f"moltbot_doc_{document.file_id}_{file_name}")
        await doc_file.download_to_drive(doc_path)

        # Processa baseado no tipo
        if file_name.endswith(".xlsx") or file_name.endswith(".xls"):
            # Excel
            import pandas as pd

            df = pd.read_excel(doc_path)

            # Limpeza e preparação profissional
            # Remove colunas completamente vazias
            df = df.dropna(axis=1, how="all")

            # Remove linhas completamente vazias
            df = df.dropna(axis=0, how="all")

            # Preenche NaN com valores apropriados
            for col in df.columns:
                if df[col].dtype == "object":
                    df[col] = df[col].fillna("-")
                else:
                    df[col] = df[col].fillna(0)

            # Identifica tipos de colunas
            numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
            text_cols = df.select_dtypes(include=["object"]).columns.tolist()
            date_cols = df.select_dtypes(include=["datetime"]).columns.tolist()

            # Prepara análise profissional
            data_summary = f"""📊 ANÁLISE DE PLANILHA EXCEL

📁 Arquivo: {file_name}
📐 Dimensões: {df.shape[0]} linhas × {df.shape[1]} colunas

📋 ESTRUTURA:
• Colunas numéricas: {", ".join(numeric_cols) if numeric_cols else "Nenhuma"}
• Colunas de texto: {", ".join(text_cols) if text_cols else "Nenhuma"}
• Colunas de data: {", ".join(date_cols) if date_cols else "Nenhuma"}

📊 AMOSTRA DOS DADOS (primeiras 8 linhas):
{df.head(8).to_markdown(index=False)}
"""

            # Adiciona estatísticas apenas se houver colunas numéricas
            if numeric_cols:
                stats = df[numeric_cols].describe().round(2)
                data_summary += (
                    f"\n\n📈 ESTATÍSTICAS (colunas numéricas):\n{stats.to_markdown()}"
                )

            # Adiciona informações de valores únicos para colunas de texto
            if text_cols and len(text_cols) <= 5:
                data_summary += "\n\n🔍 VALORES ÚNICOS (colunas de texto):"
                for col in text_cols[:3]:  # Máximo 3 colunas
                    unique_count = df[col].nunique()
                    if unique_count <= 10:
                        values = df[col].value_counts().head(5)
                        data_summary += f"\n• {col}: {values.to_dict()}"
                    else:
                        data_summary += f"\n• {col}: {unique_count} valores únicos"

            await update.message.reply_text("📊 Analisando planilha com IA...")

            prompt = f"""Você é um analista de dados profissional. Analise esta planilha e forneça um relatório executivo:

{data_summary}

FORNEÇA:
1. 📋 Resumo Executivo (2-3 frases sobre o que a planilha contém)
2. 🎯 Principais Insights (3-5 pontos importantes)
3. 📊 Análise dos Dados (padrões, tendências, anomalias)
4. 💡 Recomendações (se aplicável)

Use formatação clara com emojis e organize em seções."""

            history = []  # Sem histórico para análise limpa
            response = await agent.run(prompt, history)

            await update.message.reply_text(response[:4000])
            store.add_message("user", f"[EXCEL] {file_name}", chat_id=chat_id)
            store.add_message("assistant", response, chat_id=chat_id)

        elif file_name.endswith(".csv"):
            # CSV
            import pandas as pd

            df = pd.read_csv(doc_path)

            # Limpeza profissional
            df = df.dropna(axis=1, how="all")
            df = df.dropna(axis=0, how="all")

            for col in df.columns:
                if df[col].dtype == "object":
                    df[col] = df[col].fillna("-")
                else:
                    df[col] = df[col].fillna(0)

            numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
            text_cols = df.select_dtypes(include=["object"]).columns.tolist()

            data_summary = f"""📊 ANÁLISE DE ARQUIVO CSV

📁 Arquivo: {file_name}
📐 Dimensões: {df.shape[0]} linhas × {df.shape[1]} colunas

📋 ESTRUTURA:
• Colunas numéricas: {", ".join(numeric_cols) if numeric_cols else "Nenhuma"}
• Colunas de texto: {", ".join(text_cols) if text_cols else "Nenhuma"}

📊 AMOSTRA DOS DADOS:
{df.head(8).to_markdown(index=False)}
"""

            if numeric_cols:
                stats = df[numeric_cols].describe().round(2)
                data_summary += f"\n\n📈 ESTATÍSTICAS:\n{stats.to_markdown()}"

            await update.message.reply_text("📊 Analisando CSV com IA...")

            prompt = f"""Você é um analista de dados profissional. Analise este CSV:

{data_summary}

FORNEÇA:
1. 📋 Resumo Executivo
2. 🎯 Principais Insights
3. 📊 Análise dos Dados
4. 💡 Recomendações

Use formatação clara com emojis."""

            response = await agent.run(prompt, [])
            await update.message.reply_text(response[:4000])
            store.add_message("user", f"[CSV] {file_name}", chat_id=chat_id)
            store.add_message("assistant", response, chat_id=chat_id)

        elif file_name.endswith(".docx"):
            # Word
            from docx import Document

            doc = Document(doc_path)

            text = "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()]
            )

            preview = f"📄 **Arquivo Word:** {file_name}\n\n"
            preview += f"**Parágrafos:** {len(doc.paragraphs)}\n\n"
            preview += f"**Conteúdo:**\n{text[:3500]}"

            await update.message.reply_text(preview[:4000])
            store.add_message("user", f"[WORD] {file_name}: {len(text)} caracteres", chat_id=chat_id)

        elif file_name.endswith(".md"):
            # Markdown
            with open(doc_path, "r", encoding="utf-8") as f:
                text = f.read()

            preview = f"📝 **Arquivo Markdown:** {file_name}\n\n"
            preview += f"**Tamanho:** {len(text)} caracteres\n\n"
            preview += f"**Conteúdo:**\n{text[:3500]}"

            await update.message.reply_text(preview[:4000])
            store.add_message("user", f"[MARKDOWN] {file_name}", chat_id=chat_id)

        elif mime_type and mime_type.startswith("image/"):
            # Imagem - OCR
            from workspace.tools.extra_tools import ocr_extract

            result = await ocr_extract(doc_path)

            if result["success"]:
                text = result["text"]
                if text:
                    await update.message.reply_text(
                        f"📄 **Texto extraído (OCR):**\n\n{text[:4000]}"
                    )
                else:
                    await update.message.reply_text(
                        "⚠️ Nenhum texto encontrado na imagem."
                    )
            else:
                await update.message.reply_text("Ocorreu um erro no OCR. Tente novamente.")
        else:
            await update.message.reply_text(
                f"⚠️ Formato não suportado: {file_name}\n\n"
                "Formatos aceitos:\n"
                "• Excel (.xlsx, .xls)\n"
                "• CSV (.csv)\n"
                "• Word (.docx)\n"
                "• Markdown (.md)\n"
                "• Imagens (para OCR)"
            )

        # Limpa arquivo
        os.unlink(doc_path)

    except Exception as e:
        logger.error(f"Erro ao processar documento: {e}", exc_info=True)
        await update.message.reply_text("Ocorreu um erro ao processar o documento. Tente novamente.")
